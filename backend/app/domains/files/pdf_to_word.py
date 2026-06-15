"""PDF to Word MVP conversion helpers."""

from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from fastapi import HTTPException, status
from pypdf import PdfReader


PDF_TO_WORD_MIN_TEXT_CHARS = 40
PDF_TO_WORD_MAX_PAGES = 80


class PDFToWordError(ValueError):
    """User-facing conversion error for PDF to Word Beta."""


def convert_text_pdf_to_docx(
    *,
    input_path: str,
    output_path: str,
    insert_page_breaks: bool = True,
) -> dict:
    """Extract readable text from a text-based PDF and write a simple DOCX."""

    source = Path(input_path)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    try:
        reader = PdfReader(str(source))
    except Exception as exc:
        raise PDFToWordError("This PDF could not be opened. Try a readable text-based PDF.") from exc

    if reader.is_encrypted:
        raise PDFToWordError("Encrypted PDFs are not supported in PDF to Word Beta.")

    page_count = len(reader.pages)
    if page_count == 0:
        raise PDFToWordError("This PDF has no pages to convert.")
    if page_count > PDF_TO_WORD_MAX_PAGES:
        raise PDFToWordError(f"PDF to Word Beta supports up to {PDF_TO_WORD_MAX_PAGES} pages.")

    page_texts: list[str] = []
    total_chars = 0
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        normalized = _normalize_page_text(text)
        page_texts.append(normalized)
        total_chars += len(normalized.strip())

    if total_chars < PDF_TO_WORD_MIN_TEXT_CHARS:
        raise PDFToWordError(
            "This looks like a scanned or image-based PDF. Use OCR PDF first, then convert the recognized text."
        )

    document = Document()
    document.core_properties.title = source.stem
    document.add_heading(source.stem or "Converted PDF", level=1)

    paragraph_count = 0
    for page_index, text in enumerate(page_texts):
        if page_index > 0 and insert_page_breaks:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        if page_count > 1:
            document.add_paragraph(f"Page {page_index + 1}", style="Intense Quote")

        for block in _split_blocks(text):
            _append_block(document, block)
            paragraph_count += 1

    document.save(str(output))

    return {
        "success": True,
        "output_path": str(output),
        "file_size": os.path.getsize(output),
        "page_count": page_count,
        "paragraph_count": paragraph_count,
        "extracted_characters": total_chars,
    }


def validate_pdf_to_word_upload(file_id: str, filename: str | None) -> None:
    """Reject obviously non-PDF inputs before queueing."""

    if not file_id:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="file_id is required")

    if filename and not filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="PDF to Word Beta only accepts PDF files.",
        )


def user_facing_pdf_to_word_error(exc: Exception) -> str:
    """Return a short error message suitable for task status and history."""

    if isinstance(exc, PDFToWordError):
        return str(exc)
    message = str(exc).splitlines()[0].strip()
    return message[:240] if message else "PDF to Word conversion failed."


def _normalize_page_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_blocks(text: str) -> list[str]:
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    if blocks:
        return blocks
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return lines


def _append_block(document: Document, block: str) -> None:
    compact = " ".join(line.strip() for line in block.splitlines() if line.strip())
    if not compact:
        return

    if _looks_like_heading(compact):
        document.add_heading(compact, level=2)
        return

    document.add_paragraph(compact)


def _looks_like_heading(text: str) -> bool:
    if len(text) > 96 or text.endswith((".", ",", ";", ":")):
        return False
    if re.match(r"^(\d+(\.\d+)*|[IVXLC]+)\s+[\w(]", text):
        return True
    letters = [char for char in text if char.isalpha()]
    return bool(letters) and len(letters) >= 4 and sum(char.isupper() for char in letters) / len(letters) > 0.75
